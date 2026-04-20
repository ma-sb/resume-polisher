"""Read and parse resume files (.docx, .pdf, .doc)."""

import io
import shutil
import subprocess
import tempfile
from pathlib import Path
from dataclasses import dataclass, field
from docx import Document
from pypdf import PdfReader


@dataclass
class ResumeSection:
    heading: str
    bullets: list[str] = field(default_factory=list)


@dataclass
class Resume:
    filepath: Path
    filename: str
    full_text: str
    sections: list[ResumeSection] = field(default_factory=list)
    name: str = ""
    raw_bytes: bytes | None = None

    @property
    def display_name(self) -> str:
        return self.name or self.filename


def _is_major_heading(para) -> bool:
    """Detect top-level section headings (EXPERIENCE, EDUCATION, etc.)."""
    text = para.text.strip()
    if not text or len(text) > 80:
        return False
    if para.style.name.startswith("Heading"):
        return True
    if not para.runs:
        return False
    return all(r.bold for r in para.runs if r.text.strip())


def _parse_document(doc: Document, filename: str, filepath: Path | None = None) -> Resume:
    """Parse a python-docx Document into a Resume with per-paragraph bullets."""
    full_lines: list[str] = []
    sections: list[ResumeSection] = []
    current_section: ResumeSection | None = None
    candidate_name = ""

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        full_lines.append(text)

        if i == 0 and not candidate_name:
            candidate_name = text

        if _is_major_heading(para):
            current_section = ResumeSection(heading=text)
            sections.append(current_section)
        elif current_section is not None:
            clean = text.lstrip("•-–◦▪ ")
            current_section.bullets.append(clean)

    return Resume(
        filepath=filepath or Path(filename),
        filename=filename,
        full_text="\n".join(full_lines),
        sections=sections,
        name=candidate_name,
    )


def read_docx(filepath: Path) -> Resume:
    """Parse a .docx file from disk."""
    filepath = Path(filepath)
    raw = filepath.read_bytes()
    doc = Document(io.BytesIO(raw))
    resume = _parse_document(doc, filepath.name, filepath)
    resume.raw_bytes = raw
    return resume


def read_docx_from_bytes(data: bytes, filename: str) -> Resume:
    """Parse a .docx file from raw bytes (e.g. an uploaded file)."""
    doc = Document(io.BytesIO(data))
    resume = _parse_document(doc, filename)
    resume.raw_bytes = data
    return resume


def _convert_doc_bytes_to_docx_bytes(data: bytes, filename: str) -> bytes:
    """Convert legacy .doc bytes to .docx bytes using LibreOffice."""
    lo_cmd = None
    for candidate in ["libreoffice", "soffice", "/Applications/LibreOffice.app/Contents/MacOS/soffice"]:
        if shutil.which(candidate):
            lo_cmd = candidate
            break
    if not lo_cmd:
        raise RuntimeError("LibreOffice is required to read .doc files.")

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        in_path = tmp_path / filename
        in_path.write_bytes(data)

        result = subprocess.run(
            [lo_cmd, "--headless", "--convert-to", "docx", "--outdir", str(tmp_path), str(in_path)],
            capture_output=True,
            timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError("Failed to convert .doc file.")

        out_path = tmp_path / f"{in_path.stem}.docx"
        if not out_path.exists():
            raise RuntimeError("Converted .docx file was not produced.")
        return out_path.read_bytes()


def read_pdf_from_bytes(data: bytes, filename: str) -> Resume:
    """Parse a .pdf file from raw bytes."""
    reader = PdfReader(io.BytesIO(data))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    full_text = "\n".join(p for p in pages if p).strip()
    if not full_text:
        raise ValueError("No readable text found in PDF.")

    first_line = next((line.strip() for line in full_text.splitlines() if line.strip()), "")
    return Resume(
        filepath=Path(filename),
        filename=filename,
        full_text=full_text,
        sections=[],
        name=first_line,
        raw_bytes=None,
    )


def read_resume_from_bytes(data: bytes, filename: str) -> Resume:
    """Parse a supported resume file from raw bytes."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return read_docx_from_bytes(data, filename)
    if suffix == ".pdf":
        return read_pdf_from_bytes(data, filename)
    if suffix == ".doc":
        docx_bytes = _convert_doc_bytes_to_docx_bytes(data, filename)
        converted_name = f"{Path(filename).stem}.docx"
        return read_docx_from_bytes(docx_bytes, converted_name)
    raise ValueError(f"Unsupported file format: {suffix or 'unknown'}")


def load_resumes(folder: Path) -> list[Resume]:
    """Load all supported resumes from a folder (recursively)."""
    folder = Path(folder)
    if not folder.exists():
        return []
    resumes = []
    supported = [".docx", ".pdf", ".doc"]
    for ext in supported:
        for f in sorted(folder.rglob(f"*{ext}")):
            if f.name.startswith("~$"):
                continue
            try:
                resumes.append(read_resume_from_bytes(f.read_bytes(), f.name))
            except Exception as e:
                print(f"Warning: could not read {f.name}: {e}")
    return resumes
